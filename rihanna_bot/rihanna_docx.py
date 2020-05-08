from docx import Document
from os import path
import os
import ast
import re
from bs4 import BeautifulSoup
# https://www.youtube.com/watch?v=26vNgM_wSAE


def selector(query):
    if query[:len('word file read ')] == 'word file read ':
        f_name = query[len('word file read '):].strip()
        return WordFile(f_name).read()
    elif query[:len('word file create ')] == 'word file create ':  # word file create e.docx + content| b + heading|i +content_styled|khb
        msg = query[len('word file create '):]
        mg = msg.split('+')
        f_name = mg[0]
        dict_msg = {}
        for i in mg[1:]:
            d = i.split('|')
            dict_msg[d[0].strip()] = d[1].strip()
        return WordFile(f_name).create(**dict_msg)
    elif query[:len('word file delete ')] == 'word file delete ':
        msg = query[len('word file delete '):].strip()
        return WordFile(msg).delete()
    elif query[:len('word file send ')] == 'word file send ':
        msg = query[len('word file send '):].strip()
        q = ast.literal_eval(msg)
        return WordFile(q['filename']).create(html=q['content'])
    else:
        reply = 'i do not know'
        return {'display': reply, 'say': reply}


class WordFile:
    def __init__(self, filename):
        self.path = r'C:\Users\emyli\PycharmProjects\Chatbot_Project\wordfiles'
        self.filename = filename

    def create(self, content=None, content_styled=None, heading=None, html=None):
        # content = just text, content_styled = [[{'content':'i ..', 'style':'bold' or'italic' or None]},{}...], ..]
        if path.exists(f'{self.path}\{self.filename}'):
            reply = f'{self.filename} exists in directory'
            return {'display': reply, 'say': reply}
        doc = Document()
        if heading:
            doc.add_heading(heading, 0)
        if content:
            doc.add_paragraph(content)
            doc.save(f'{self.path}\{self.filename}')
            reply = f'Word document {self.filename} has been created'
            return {'display': reply, 'say': reply}
        elif content_styled:
            if type(content_styled).__name__ == 'str':
                try:
                    content_styled = ast.literal_eval(content_styled)
                except:
                    reply = 'improper formatting on WordFile'
                    return {'display': reply, 'say': reply}
            for para in content_styled:
                p = doc.add_paragraph('')
                for run in para:
                    if run['style'] == 'bold':
                        p.add_run(run['content']).bold = True
                    elif run['style'] == 'italic':
                        p.add_run(run['content']).italic = True
                    else:
                        p.add_run(run['content'])
            doc.save(f'{self.path}\{self.filename}')
            reply = f'Word document {self.filename} has been created'
            return {'display': reply, 'say': reply}
        elif html:
            b = html.split('\n')
            for i in b:
                soup = BeautifulSoup(i, 'html.parser')
                no = re.findall('<+[a-z0-9]+>', i)
                if len(no) == 0:
                    doc.add_paragraph(i)
                elif len(no) == 1:
                    for tag in no:
                        t = re.findall('\w+', tag)[0]
                        print(t)
                        if t == 'h1':
                            # soup = BeautifulSoup(i, 'html.parser')
                            title = soup.find('h1').text
                            doc.add_heading(title, 0)
                        elif t[0] == 'h':
                            # soup = BeautifulSoup(i, 'html.parser')
                            data = soup.find(t).text
                            doc.add_heading(data)
                        else:
                            p = doc.add_paragraph('')
                            data = soup.find(t).text
                            if t == 'b':
                                p.add_run(data).bold = True
                            elif t == 'em':
                                p.add_run(data).italic = True
                else:
                    p = doc.add_paragraph('')
                    for tag in no:
                        t = re.findall('\w+', tag)
                        for wa in t:
                            data = soup.find(wa).text
                            if t == 'b':
                                p.add_run(data).bold = True
                            elif t == 'em':
                                p.add_run(data).italic = True
            doc.save(f'{self.path}\{self.filename}')
            reply = f'Word document {self.filename} has been created'
            return {'display': reply, 'say': reply}

        else:
            reply = 'no content specified'
            return {'display': reply, 'say': reply}

    def read(self):
        if path.exists(f'{self.path}\{self.filename}'):
            display_content = ''
            say_content = ''
            doc = Document(f'{self.path}\{self.filename}')
            for para in doc.paragraphs:
                #print(para.font)
                if para.style.name == 'Title':
                    display_content += f'<h1>{para.text}</h1>'
                    say_content += para.text + '\n'
                else:
                    say_content += para.text + '\n'
                    p = '<p>'
                    for run in para.runs:
                        if run.font.bold:
                            p += f'<b>{run.text}</b>'
                        elif run.font.italic:
                            p += f'<em>{run.text}</em>'
                        else:
                            p += f'{run.text}'
                    p += '</p>'
                    display_content += p
            return {'display': display_content, 'say': say_content}

        else:
            reply = f'file {self.filename} does not exist'
            return {'display': reply, 'say': reply}

    def delete(self):
        if path.exists(f'{self.path}\{self.filename}'):
            os.remove(f'{self.path}\{self.filename}')
            reply = f'Word document {self.filename} has been deleted'
            return {'display': reply, 'say': reply}
        else:
            reply = f'file {self.filename} does not exist'
            return {'display': reply, 'say': reply}


# WordFile('emeka.docx').delete()
# con = [
#         [{'content':'i love people', 'style':'bold'}, {'content':' they are very nice', 'style':'italic'}],
#         [{'content':'i love people', 'style':'italic'}, {'content':' they are very nice', 'style':None}],
#        [{'content': 'i love people', 'style': 'bold'}, {'content': ' they are very nice', 'style': 'bold'}],
#        [{'content': 'i love people', 'style': None}, {'content': ' they are very nice', 'style': 'italic'}]
#        ]
# a = WordFile('emeka.docx').create(content_styled=con, heading='life is good')
#
# print(a)

# a = WordFile('emeka.docx').read()
# print(a)

# d = "word file create file.docx + heading| Test + " \
#     "content_styled|[[{'content':'i love people', 'style':'bold'}, " \
#     "{'content':' they are very nice', 'style':'italic'}]]"
# a = selector(d)
#
# print(a)


# import re
# from bs4 import BeautifulSoup
#
# a = ' i am good \n i love you \n <b>thank you</b> \n my tahns <b>i knwp</b>'
# b = a.split('\n')
# con = []
# for i in b:
#     no = re.findall('<+[a-z]+>', i)
#     if len(no) == 0:
#         con.append([{'content': i, 'style':None}])
#     else:
#         style = {'b': 'bold', 'em': 'italic'}
#         para = []
#         for tag in no:
#             t = re.findall('\w+', tag)
# a = '<h1> i am good</h1> \n i <b>love</b> you \n <b>thank you</b> \n my tahns <b>i knwp</b>'
# WordFile('tesr.docx').create(html=a)